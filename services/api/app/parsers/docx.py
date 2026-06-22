from pathlib import Path

from docx import Document as DocxDocument

from app.parsers.interface import ParsedBlock, ParsedDocument, ParsedPage
from app.parsers.sections import detect_section


def parse_docx_file(file_path: str) -> ParsedDocument:
    path = Path(file_path)
    document = DocxDocument(file_path)
    blocks: list[ParsedBlock] = []
    current_section: str | None = None

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        is_heading = style_name.lower().startswith("heading")
        if is_heading:
            current_section = detect_section(text) or text
            blocks.append(
                ParsedBlock(
                    block_type="heading",
                    page_number=1,
                    section=current_section,
                    text=text,
                    metadata={"style": style_name},
                )
            )
            continue

        blocks.append(ParsedBlock(block_type="paragraph", page_number=1, section=current_section, text=text))

    page_text = "\n\n".join(block.text or "" for block in blocks)
    return ParsedDocument(
        metadata={"parser": "docx", "filename": path.name},
        pages=[ParsedPage(page_number=1, text=page_text)],
        blocks=blocks,
    )
